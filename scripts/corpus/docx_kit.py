"""Shared python-docx scaffolding."""
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# Fixed so regenerating does not churn the document metadata.
FIXED_TIMESTAMP = datetime(2026, 1, 15, 9, 0, 0)


def new_document(title: str, subtitle: str) -> Document:
    document = Document()
    document.core_properties.title = title
    document.core_properties.author = "Aurora Home"
    document.core_properties.created = FIXED_TIMESTAMP
    document.core_properties.modified = FIXED_TIMESTAMP

    document.add_heading(title, level=0)
    lead = document.add_paragraph(subtitle)
    lead.runs[0].italic = True
    lead.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return document


def add_table(document: Document, header: list[str], rows: list[list[str]]):
    """Ruled table — `Table Grid` gives every cell a visible border."""
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, header, strict=True):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row, strict=True):
            cell.text = text
    return table


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_figure(document: Document, image: Path, caption: str | None) -> None:
    document.add_picture(str(image), width=Inches(5.4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption is not None:
        add_caption(document, caption)


def _normalise_zip(path: Path) -> None:
    """Rewrite the .docx with fixed zip timestamps and a stable entry order.

    A .docx is a zip, and python-docx stamps each entry with the time of writing
    — so two runs produce different bytes even with identical content. Content
    hashes drive the incremental-ingestion manifest, so that churn matters.
    """
    fixed = (FIXED_TIMESTAMP.year, FIXED_TIMESTAMP.month, FIXED_TIMESTAMP.day, 9, 0, 0)
    temporary = path.with_suffix(path.suffix + ".tmp")

    with zipfile.ZipFile(path) as source:
        entries = sorted(source.infolist(), key=lambda item: item.filename)
        payload = [(item, source.read(item.filename)) for item in entries]

    with zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for item, data in payload:
            info = zipfile.ZipInfo(item.filename, date_time=fixed)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = item.external_attr
            target.writestr(info, data)

    temporary.replace(path)


def save(document: Document, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    _normalise_zip(path)
    return path
