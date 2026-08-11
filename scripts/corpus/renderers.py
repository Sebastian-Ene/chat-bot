"""Render a parsed Markdown `Doc` to PDF, DOCX or HTML.

`layout` in the front matter selects a visual treatment. The variants exist so
the bulk of the corpus does not share one structural fingerprint — a uniform
corpus quietly flatters the parser.
"""
from pathlib import Path

from docx.shared import Pt
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from reportlab.platypus import ListFlowable, ListItem, NextPageTemplate, PageBreak, Paragraph, Spacer

from scripts.corpus import docx_kit, markdown_kit, paths, pdf_kit


def _figures(blocks: list[markdown_kit.Block]) -> list[str]:
    """Figure references anywhere in the document, including inside columns."""
    found: list[str] = []
    for block in blocks:
        if block.kind == "figure":
            found.append(block.text)
        elif block.kind == "columns":
            found += _figures(block.blocks)
    return found

# font, body alignment, body size, leading, heading scale
PDF_LAYOUTS = {
    "report": ("Helvetica", TA_JUSTIFY, 10.0, 14.5, 1.0),
    "manual": ("Times-Roman", TA_LEFT, 10.5, 15.0, 1.08),
    "brief": ("Helvetica", TA_LEFT, 9.5, 13.0, 0.92),
}

HTML_LAYOUTS = {
    "report": ("Georgia, serif", "52rem", "#1b1b1b"),
    "manual": ("'Helvetica Neue', Arial, sans-serif", "46rem", "#222222"),
    "brief": ("'Iowan Old Style', Georgia, serif", "40rem", "#2b2b2b"),
}


# --------------------------------------------------------------------- PDF ---
def _pdf_styles(layout: str) -> dict:
    style = pdf_kit.styles()
    font, alignment, size, leading, scale = PDF_LAYOUTS.get(layout, PDF_LAYOUTS["report"])
    style["body"].fontName = font
    style["body"].alignment = alignment
    style["body"].fontSize = size
    style["body"].leading = leading
    style["bullet"].fontName = font
    for name, base in (("h1", 15), ("h2", 12.5)):
        style[name].fontSize = base * scale
        style[name].leading = base * scale * 1.25
    return style


def _pdf_blocks(blocks: list[markdown_kit.Block], style: dict, width: float) -> list:
    story: list = []
    for block in blocks:
        if block.kind == "heading":
            story.append(Paragraph(block.text, style["h1" if block.level <= 1 else "h2"]))
        elif block.kind == "para":
            story.append(Paragraph(block.text, style["body"]))
        elif block.kind == "bullets":
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(item, style["bullet"])) for item in block.items],
                    bulletType="bullet",
                    start="•",
                )
            )
            story.append(Spacer(1, 6))
        elif block.kind == "table":
            data = [block.header, *block.rows]
            widths = [width / len(block.header)] * len(block.header)
            table = pdf_kit.ruled_table(data, widths) if block.ruled else pdf_kit.unruled_table(data, widths)
            story.append(table)
            story.append(Spacer(1, 10))
        elif block.kind == "figure":
            story += pdf_kit.figure(paths.image_source(block.text), width * 0.8, block.caption, style["caption"])
        elif block.kind == "columns":
            # Column changes apply per page in ReportLab, so the block is fenced
            # by page breaks.
            story.append(NextPageTemplate("TwoCol"))
            story.append(PageBreak())
            story += _pdf_blocks(block.blocks, style, width / 2)
            story.append(NextPageTemplate("OneCol"))
            story.append(PageBreak())
    return story


def render_pdf(doc: markdown_kit.Doc, path: Path) -> Path:
    def story(_default_styles: dict, usable: float) -> list:
        style = _pdf_styles(doc.layout)
        head = [Paragraph(doc.title, style["title"])]
        if doc.meta.get("subtitle"):
            head.append(Paragraph(doc.meta["subtitle"], style["subtitle"]))
        return head + _pdf_blocks(doc.blocks, style, usable)

    return pdf_kit.build_pdf(path, doc.title, "Aurora Home", story)


# -------------------------------------------------------------------- DOCX ---
def _docx_blocks(document, blocks: list[markdown_kit.Block]) -> None:
    for block in blocks:
        if block.kind == "heading":
            document.add_heading(block.text, level=min(block.level, 4))
        elif block.kind == "para":
            document.add_paragraph(block.text)
        elif block.kind == "bullets":
            for item in block.items:
                document.add_paragraph(item, style="List Bullet")
        elif block.kind == "table":
            table = docx_kit.add_table(document, block.header, block.rows)
            if not block.ruled:
                table.style = "Normal Table"
        elif block.kind == "figure":
            docx_kit.add_figure(document, paths.image_source(block.text), block.caption)
        elif block.kind == "columns":
            # No column support here: DOCX columns are a section property, and a
            # mid-document section break is not worth the complexity.
            _docx_blocks(document, block.blocks)


def render_docx(doc: markdown_kit.Doc, path: Path) -> Path:
    document = docx_kit.new_document(doc.title, doc.meta.get("subtitle", ""))
    if doc.layout == "manual":
        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
    elif doc.layout == "brief":
        document.styles["Normal"].font.size = Pt(9.5)
    _docx_blocks(document, doc.blocks)
    return docx_kit.save(document, path)


# -------------------------------------------------------------------- HTML ---
def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _html_blocks(blocks: list[markdown_kit.Block]) -> list[str]:
    out: list[str] = []
    for block in blocks:
        if block.kind == "heading":
            level = min(block.level + 1, 4)
            out.append(f"<h{level}>{_escape(block.text)}</h{level}>")
        elif block.kind == "para":
            out.append(f"<p>{_escape(block.text)}</p>")
        elif block.kind == "bullets":
            items = "".join(f"<li>{_escape(item)}</li>" for item in block.items)
            out.append(f"<ul>{items}</ul>")
        elif block.kind == "table":
            head = "".join(f"<th>{_escape(cell)}</th>" for cell in block.header)
            body = "".join(
                "<tr>" + "".join(f"<td>{_escape(cell)}</td>" for cell in row) + "</tr>"
                for row in block.rows
            )
            css = "" if block.ruled else ' class="unruled"'
            out.append(f"<table{css}><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>")
        elif block.kind == "figure":
            caption = f"<figcaption>{_escape(block.caption)}</figcaption>" if block.caption else ""
            out.append(f'<figure><img src="{block.text}" alt="" />{caption}</figure>')
        elif block.kind == "columns":
            out.append('<div class="two-column">')
            out += _html_blocks(block.blocks)
            out.append("</div>")
    return out


def render_html(doc: markdown_kit.Doc, path: Path) -> Path:
    font, measure, ink = HTML_LAYOUTS.get(doc.layout, HTML_LAYOUTS["report"])
    subtitle = f'<p class="lede">{_escape(doc.meta["subtitle"])}</p>' if doc.meta.get("subtitle") else ""
    body = "\n    ".join(_html_blocks(doc.blocks))
    html = f"""<!doctype html>
<html lang="{doc.lang}">
  <head>
    <meta charset="utf-8" />
    <title>{_escape(doc.title)}</title>
    <style>
      body {{ font-family: {font}; max-width: {measure}; margin: 2rem auto; color: {ink}; }}
      h1 {{ font-size: 1.8rem; }}
      h2 {{ font-size: 1.25rem; margin-top: 1.8rem; border-bottom: 1px solid #ddd; }}
      h3 {{ font-size: 1.05rem; }}
      .lede {{ color: #555; font-style: italic; }}
      .two-column {{ column-count: 2; column-gap: 2.5rem; }}
      figure img {{ width: 100%; max-width: 32rem; }}
      table {{ border-collapse: collapse; margin: 1rem 0; }}
      th, td {{ border: 1px solid #999; padding: .3rem .55rem; text-align: left; }}
      th {{ background: #e8e3d8; }}
      table.unruled th, table.unruled td {{ border: none; padding: .2rem 1.2rem .2rem 0; }}
    </style>
  </head>
  <body>
    <h1>{_escape(doc.title)}</h1>
    {subtitle}
    {body}
  </body>
</html>
"""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(html, encoding="utf-8")
    # Only HTML needs its figures on disk beside it; PDF and DOCX embed them.
    paths.copy_images(_figures(doc.blocks), path.parent)
    return path


RENDERERS = {"pdf": render_pdf, "docx": render_docx, "html": render_html}


def render(doc: markdown_kit.Doc, out_dir: Path, stem: str) -> Path:
    renderer = RENDERERS[doc.fmt]
    return renderer(doc, out_dir / f"{stem}.{doc.fmt}")
